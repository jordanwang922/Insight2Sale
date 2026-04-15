#!/usr/bin/env python3
"""生成《Insight2Sale 客户自部署手册》Word 文档。依赖: python-docx"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def main() -> None:
    out = Path(__file__).resolve().parent.parent / "docs" / "client-deployment"
    out.mkdir(parents=True, exist_ok=True)
    path = out / "Insight2Sale-客户自部署手册-v0.8.docx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Insight2Sale 客户自部署手册")
    r.bold = True
    r.font.size = Pt(18)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("（v0.8 基线 · 含服务器、HTTPS、AI 密钥与数据迁移说明）")
    sr.font.size = Pt(11)

    doc.add_paragraph()

    add_heading(doc, "文档说明", 1)
    add_para(
        doc,
        "本文面向需要在自有服务器（含阿里云 ECS 等）上部署 Insight2Sale 的交付方或客户运维人员。"
        "内容整合：服务器规格与软件、从 GitHub 拉取代码、数据库与迁移、环境变量（含火山引擎方舟与豆包语音）、"
        "HTTPS、进程守护、验证步骤，以及演示数据与知识库迁移注意事项。",
    )

    add_heading(doc, "一、源代码与 GitHub 访问方式", 1)
    add_para(doc, "官方仓库地址（请核对是否为最新）：", bold=True)
    add_para(doc, "https://github.com/jordanwang922/Insight2Sale")
    add_heading(doc, "1.1 公开仓库与私有仓库的选择", 2)
    add_bullets(
        doc,
        [
            "若仓库为 Public：任何知道链接的人均可 git clone 源码。请避免在仓库或 Issue 中提交 .env、密钥与真实客户数据。",
            "若仓库为 Private：客户无法用匿名方式拉代码。常见做法包括：在 GitHub 将客户指定的 GitHub 账号添加为 Collaborator（建议只读）、"
            "或为部署机配置 Deploy SSH Key / Fine-grained Personal Access Token（仅勾选 Contents: Read）。",
            "交付物中的「程序」以 Git 仓库为准；数据库与 storage 中的上传文件需单独迁移（见后文）。",
        ],
    )
    add_heading(doc, "1.2 获取代码（在服务器上执行）", 2)
    add_para(doc, "建议使用固定版本标签部署，便于回滚与对齐文档：", bold=True)
    add_para(doc, "git clone https://github.com/jordanwang922/Insight2Sale.git")
    add_para(doc, "cd Insight2Sale")
    add_para(doc, "git fetch --tags")
    add_para(doc, "git checkout v0.8.0   # 或与交付方约定的 tag / commit")
    add_para(doc, "说明：若使用 SSH 地址，请将 HTTPS URL 换为 git@github.com:jordanwang922/Insight2Sale.git，并先在服务器配置 SSH 公钥。", bold=True)

    add_heading(doc, "二、服务器与运行环境建议", 1)
    add_heading(doc, "2.1 硬件与操作系统", 2)
    add_bullets(
        doc,
        [
            "计算：建议 2 核 CPU / 4 GB 内存起步；若知识库 PDF OCR、并发用户较多，建议 4 核 8 GB 或以上。",
            "磁盘：系统盘 + 数据盘规划充足空间；知识库文件与通话录音会持续增长。",
            "操作系统：Alibaba Cloud Linux 3、Ubuntu 22.04 LTS 等与 Node 20+ 兼容的 64 位系统。",
        ],
    )
    add_heading(doc, "2.2 需要预先准备或开通的内容", 2)
    add_bullets(
        doc,
        [
            "公网 IP 与域名（推荐）：用于 HTTPS 与浏览器访问；备案按当地法规执行。",
            "安全组：仅开放 22（建议限制来源 IP）、80、443；勿对公网开放数据库端口。",
            "PostgreSQL：推荐使用阿里云 RDS PostgreSQL，或与 ECS 同机/内网自建 PostgreSQL（仅内网访问）。",
            "可选：对象存储 OSS（当前版本默认使用服务器本地 storage 目录，迁移时需整目录备份）。",
        ],
    )
    add_heading(doc, "2.3 服务器软件清单", 2)
    add_bullets(
        doc,
        [
            "Node.js：建议 LTS v20 或以上（与构建一致）。",
            "npm：随 Node 安装；部署时使用 npm ci 或 npm install。",
            "Nginx：反向代理、HTTPS、上传大小限制。",
            "进程管理：PM2 或 systemd，保证进程崩溃自拉起与开机启动。",
            "Git：用于拉取代码。",
            "可选：Certbot 或阿里云 SSL 证书，用于 HTTPS。",
        ],
    )

    add_heading(doc, "三、需要向云厂商与 AI 服务申请的内容", 1)
    add_heading(doc, "3.1 火山引擎：方舟大模型（文本 / JSON，与「豆包语音」不是同一套密钥）", 2)
    add_bullets(
        doc,
        [
            "用途：客户解读台相关 AI、通话录音转写后的纪要/要点、测评表 AI 草案等（代码通过 ARK_* 调用）。",
            "需在火山方舟控制台创建 API Key，并创建模型推理接入点；将接入点名称填入环境变量 ARK_MODEL。",
            "环境变量：ARK_API_KEY、ARK_BASE_URL（一般为 https://ark.cn-beijing.volces.com/api/v3，以控制台为准）、ARK_MODEL、ARK_TIMEOUT_MS。",
        ],
    )
    add_heading(doc, "3.2 火山引擎：豆包语音（通话录音转文字，妙记式分句）", 2)
    add_bullets(
        doc,
        [
            "用途：上传录音后的语音识别；与方舟密钥分离。",
            "需在「豆包语音」控制台创建应用，取得 AppKey 与 Access Key；开通录音文件识别相关能力。",
            "环境变量：VOLC_SPEECH_APP_KEY、VOLC_SPEECH_ACCESS_KEY；可选 VOLC_SPEECH_RESOURCE_ID（默认 volc.bigasr.auc_turbo）。",
            "若未配置豆包语音，系统可尝试使用 OpenAI Whisper 作为备选（需 OPENAI_API_KEY），生产环境建议优先豆包语音。",
        ],
    )
    add_heading(doc, "3.3 可选：OpenAI", 2)
    add_para(doc, "OPENAI_API_KEY：仅在未配置豆包语音时作为录音转写备选；非必须。")

    add_heading(doc, "四、环境变量与 .env 配置", 1)
    add_para(doc, "在项目根目录复制 .env.example 为 .env，权限建议 chmod 600。以下为要点（具体以交付时 .env.example 为准）：", bold=True)
    tbl = doc.add_table(rows=1, cols=3)
    hdr = tbl.rows[0].cells
    hdr[0].text = "变量"
    hdr[1].text = "说明"
    hdr[2].text = "是否必填"
    rows = [
        ("DATABASE_URL", "PostgreSQL 连接串", "必填"),
        ("AUTH_SECRET", "随机长字符串，用于会话加密", "必填"),
        ("AUTH_URL", "对外访问的根 URL，须含 https 与域名，与浏览器地址一致", "必填"),
        ("ARK_API_KEY / ARK_BASE_URL / ARK_MODEL / ARK_TIMEOUT_MS", "方舟大模型", "功能需要时必填"),
        ("VOLC_SPEECH_APP_KEY / VOLC_SPEECH_ACCESS_KEY", "豆包语音转写", "强烈建议"),
        ("VOLC_SPEECH_RESOURCE_ID", "语音资源 ID", "可选"),
        ("OPENAI_API_KEY", "Whisper 备选转写", "可选"),
    ]
    for a, b, c in rows:
        row = tbl.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c
    doc.add_paragraph()
    add_para(doc, "PDF 知识库 OCR 相关变量（PDF_OCR_*）见 .env.example 注释；知识库向量为本地 local-hash-v1，不依赖外部 Embedding API。", bold=True)

    add_heading(doc, "五、数据库：迁移、演示数据与知识库整体搬迁", 1)
    add_heading(doc, "5.1 全新空库部署", 2)
    add_numbered(
        doc,
        [
            "创建空数据库并配置 DATABASE_URL。",
            "在项目目录执行：npx prisma migrate deploy",
            "若需演示账号：npm run db:seed（会创建主管与两名销售等演示数据，详见 README）。",
        ],
    )
    add_heading(doc, "5.2 保留交付方已上传的知识库（推荐）", 2)
    add_para(
        doc,
        "知识库条目与向量存在 PostgreSQL 中；原始文件路径记录在库中，文件本体位于服务器 storage 目录（具体子目录以运行时为准）。"
        "仅恢复数据库而不同步 storage 会导致「记录存在但文件打不开」。建议交付方提供：",
        bold=True,
    )
    add_bullets(
        doc,
        [
            "数据库备份：使用 pg_dump（自定义格式或 SQL），在客户环境恢复。",
            "文件备份：将交付环境整个 storage 目录（或至少知识库相关上传目录）打包，在客户服务器相同项目路径下解压覆盖。",
            "恢复顺序：先恢复文件与目录权限，再导入数据库，最后执行 prisma migrate deploy 确认 schema 版本与迁移一致（若备份来自已对齐的 schema，按交付方说明操作）。",
        ],
    )
    add_para(doc, "具体 pg_dump / pg_restore 命令需与交付方约定数据库名、用户与路径；生产环境务必先备份客户侧再操作。", bold=True)

    add_heading(doc, "5.3 演示账号与测试数据说明", 2)
    add_bullets(
        doc,
        [
            "种子数据（npm run db:seed）默认创建一名主管与两名销售，默认密码见仓库 README；首次登录后请立即修改密码。",
            "主管账号可在后台使用「新增销售」创建更多销售账号；当前系统未提供「删除销售账号」界面，"
            "若需删除多余测试销售，请在数据库中由运维人员操作 User 表（注意外键与客户归属），或向交付方索要维护脚本。",
            "重要：当前版本应用内没有「删除客户」功能；演示客户、测评记录若需清理，需在 PostgreSQL 中由 DBA 执行 SQL（注意备份与合规）。",
            "若您通过 pg_dump 整体交付含知识库的库：可同时保留知识库数据；演示客户是否保留由业务决定，与知识库表无必然绑定。",
        ],
    )

    add_heading(doc, "六、构建与启动", 1)
    add_numbered(
        doc,
            [
                "cd Insight2Sale",
                "npm ci",
                "export NODE_ENV=production",
                "npm run build",
                "npm run start -- --hostname 127.0.0.1 --port 3000   # 或约定端口，与 Nginx 一致",
            ],
        )
    add_para(doc, "生产环境建议使用 PM2 或 systemd 托管上述 npm run start，并设置开机自启。", bold=True)

    add_heading(doc, "七、Nginx 反向代理与 HTTPS", 1)
    add_bullets(
        doc,
        [
            "将 server_name 指向您的域名；proxy_pass 到 127.0.0.1:应用端口。",
            "设置 proxy_set_header Host、X-Forwarded-Proto、X-Forwarded-For，以便 NextAuth 与 Cookie 在 HTTPS 下正常工作。",
            "设置 client_max_body_size（例如 50m），避免知识库或录音上传被 Nginx 413 拒绝。",
            "配置 SSL：可使用 Let's Encrypt（Certbot）或阿里云下载的证书文件。",
            "将 AUTH_URL 设为 https://您的域名，与证书域名一致。",
        ],
    )

    add_heading(doc, "八、上线验证清单", 1)
    add_bullets(
        doc,
        [
            "浏览器可打开 https 域名并进入登录页。",
            "使用主管或销售账号登录成功。",
            "随机打开客户解读台、知识库列表（若已迁移数据应能看到条目）。",
            "上传小文件测试知识库（可选）。",
            "在宽屏解读台测试通话录音上传与「通话管理」列表（需配置豆包语音密钥）。",
            "检查服务器磁盘与 CPU 在 OCR 或大文件上传时是否充足。",
        ],
    )

    add_heading(doc, "九、常见问题", 1)
    add_bullets(
        doc,
        [
            "登录异常：检查 AUTH_URL 是否与浏览器地址完全一致（含 https、域名、无多余端口）。",
            "数据库报错缺列缺表：未执行 prisma migrate deploy 或备份与代码版本不一致。",
            "转写失败：检查 VOLC_SPEECH_* 与火山控制台开通情况；或临时配置 OPENAI_API_KEY。",
            "知识库有记录但文件 404：storage 目录未同步或路径不一致。",
            "本机 3000 端口被占用：修改 next start 端口与 Nginx upstream 一致。",
        ],
    )

    add_heading(doc, "十、支持与文档索引", 1)
    add_bullets(
        doc,
        [
            "产品设计基线：仓库内 docs/design/system-design.md",
            "交接说明：docs/logs/handoff-log.md",
            "README.md：本地启动与默认账号摘要",
        ],
    )

    doc.save(path)
    print(f"Wrote: {path}")


if __name__ == "__main__":
    main()
