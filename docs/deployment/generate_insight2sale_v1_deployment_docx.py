#!/usr/bin/env python3
"""生成 Insight2Sale v1.0 客户版部署指南 Word 文档（一次性脚本，可重复运行覆盖同名文件）。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_cell_shading(cell, fill_hex: str) -> None:
    """浅灰表头背景。"""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(11)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    p.paragraph_format.space_after = Pt(6)


def main() -> None:
    out = Path(__file__).resolve().parent / "Insight2Sale-v1.0-部署指南-客户版.docx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Insight2Sale 销售型 CRM 系统\nv1.0 生产环境部署指南（客户执行版）")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        "本文面向贵司运维 / 技术人员，按步骤 1、2、3… 说明从空服务器到可访问系统的全流程。\n"
        "包含：操作系统与依赖、数据库、火山方舟（对话 + 知识库向量）、可选豆包语音、环境变量、构建启动、HTTPS、持久化与验收。"
    )
    sr.font.size = Pt(10.5)
    sr.font.name = "Microsoft YaHei"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    doc.add_paragraph()

    add_heading(doc, "一、文档说明与适用对象", 1)
    add_para(
        doc,
        "1. 适用对象：贵司已提供一台 Linux 服务器 root 权限（或等价管理员权限），由贵司技术人员按本文完成部署；"
        "若由我方远程协助，仍建议以本文为检查清单。",
    )
    add_para(
        doc,
        "2. 软件版本：本文对应 Insight2Sale 发行版本 v1.0.0（与代码仓库 tag v1.0.0、package.json 版本一致）。",
    )
    add_para(
        doc,
        "3. 外部服务：系统依赖「火山引擎 · 方舟大模型」用于对话与知识库语义向量；"
        "通话录音转写另可选用「豆包语音」或备选 OpenAI Whisper。均需在对应控制台开通并获取密钥。",
    )
    add_para(
        doc,
        "4. 重要说明：知识库「向量」存于贵司自有 PostgreSQL 数据库字段中，通过方舟 API 在线生成向量；"
        "无需单独购买火山「托管向量数据库」独立产品，但必须在方舟控制台为 Embedding 单独创建推理接入点。",
    )

    add_heading(doc, "二、建议服务器与网络条件", 1)
    add_para(doc, "在开始之前，请确认以下条件已具备或已规划：")
    items = [
        "操作系统：推荐 Ubuntu Server 22.04 LTS 或同类长期支持版本（64 位）。",
        "CPU / 内存：建议至少 2 vCPU、4GB RAM（知识库 PDF 与并发略多时建议 8GB）。",
        "磁盘：系统盘 + 数据盘规划；需为 PostgreSQL 数据、上传文件、通话录音预留空间（录音与上传随业务增长）。",
        "网络：服务器可访问公网，用于调用火山方舟 API；浏览器用户通过 HTTPS 访问贵司域名。",
        "域名与证书：生产环境强烈建议使用 HTTPS（Let's Encrypt 或贵司证书）；下文给出 Nginx 反向代理思路。",
    ]
    for t in items:
        p = doc.add_paragraph(t, style="List Bullet")
        for run in p.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    add_heading(doc, "三、步骤总览（您要依次完成的事）", 1)
    overview = [
        "准备服务器与基础环境（账户、防火墙、时区）。",
        "安装 Node.js、包管理器、Git；安装并启动 PostgreSQL（可用 Docker 或系统包）。",
        "在火山引擎开通方舟：对话模型接入点 + 知识库 Embedding 接入点（两个不同接入点 ID）。",
        "（可选）开通豆包语音，用于通话录音转写。",
        "获取代码、配置 .env、执行数据库迁移与（可选）种子数据。",
        "若启用知识库向量：对已入库文档执行一次向量重算（或上线后首次上传前配置好 Embedding）。",
        "构建前端、配置进程守护（如 PM2）与反向代理（Nginx）+ HTTPS。",
        "配置录音与上传目录持久化，完成验收测试。",
        "（可选）若需扫描版 PDF 文字识别，在服务器安装 Tesseract 并按需配置 PDF_OCR_*。",
    ]
    for i, t in enumerate(overview, 1):
        add_para(doc, f"{i}. {t}")

    add_heading(doc, "四、步骤 1：服务器基础准备", 1)
    add_para(
        doc,
        "使用 root 或 sudo 用户登录服务器。建议：设置时区（如 Asia/Shanghai）、配置防火墙仅开放 22（SSH）、80、443，"
        "数据库端口 5432 建议仅本机访问勿对公网暴露。",
    )
    add_para(doc, "创建运行用户（可选，推荐非 root 跑 Node 进程）：例如新增用户 insight2sale，后续项目目录归该用户所有。")

    add_heading(doc, "五、步骤 2：安装 Node.js 与 Git", 1)
    add_para(
        doc,
        "本项目为 Next.js 16 应用，建议使用 Node.js 20 LTS 或官方当前 LTS 版本。"
        "可使用 nvm 安装，或从 NodeSource 等渠道安装系统级 Node。同时安装 Git 以便拉取代码。",
    )
    add_para(
        doc,
        "验证命令示例：node -v（应显示 v20.x 或兼容版本）；npm -v；git --version。",
    )

    add_heading(doc, "六、步骤 3：安装 PostgreSQL", 1)
    add_para(
        doc,
        "方式 A：使用项目自带 compose.yaml 在本机启动 PostgreSQL 16 容器（需已安装 Docker 与 Docker Compose）。"
        "默认库名 insight2sale、用户 postgres、密码 postgres，端口映射 5432。生产环境请修改默认密码并限制监听地址。",
    )
    add_para(
        doc,
        "方式 B：使用系统包管理器安装 PostgreSQL 16，创建数据库与用户，记下 DATABASE_URL，格式示例：",
    )
    add_para(
        doc,
        "postgresql://用户名:密码@127.0.0.1:5432/insight2sale?schema=public",
    )

    add_heading(doc, "七、步骤 4：火山引擎方舟 — 必须开通的两类能力", 1)
    add_heading(doc, "4.1 对话大模型（与解读台、纪要等共用）", 2)
    add_para(
        doc,
        "登录火山引擎控制台，进入「机器学习平台 / 方舟」或当前产品入口，创建在线推理接入点，选择贵司采购的豆包对话模型。"
        "记录：API Key、Base URL（一般为 https://ark.cn-beijing.volces.com/api/v3 或控制台所示）、推理接入点 ID（将填入环境变量 ARK_MODEL）。",
    )
    add_heading(doc, "4.2 知识库语义向量（Embedding，与对话模型不是同一个接入点）", 2)
    add_para(
        doc,
        "在方舟控制台再创建一个专门用于「文本向量 / Embedding」的推理接入点，模型类型选择 Doubao-embedding 或控制台提供的文本向量模型，"
        "记录该接入点 ID（形如 ep-xxxx，将填入 ARK_EMBEDDING_MODEL）。切勿与 ARK_MODEL 混用。",
    )
    add_para(
        doc,
        "若控制台提供的是「多模态 / Vision」类 Embedding（例如 Doubao-embedding-vision），需在环境变量中设置 ARK_EMBEDDING_USE_MULTIMODAL=1，"
        "系统将走 /embeddings/multimodal 接口；纯文本 Embedding 接入点则不要设置或设为 0。",
    )
    add_para(
        doc,
        "说明：向量数据存储在贵司 PostgreSQL 表 KnowledgeChunk 中，不依赖火山单独「向量数据库」SaaS；"
        "但必须开通上述 Embedding 推理接入点，否则知识库只能降级为开发用本地哈希模式，检索质量不符合生产预期。",
    )

    add_heading(doc, "八、步骤 5（可选）：豆包语音 — 通话录音转写", 1)
    add_para(
        doc,
        "若使用「解读台录音 → 妙记式转写」功能，需在火山引擎「豆包语音」或控制台指引处单独开通应用，获取 AppKey、Access Key 等，"
        "填入 VOLC_SPEECH_APP_KEY、VOLC_SPEECH_ACCESS_KEY（详见 .env.example）。未配置时可降级为 OpenAI Whisper（需 OPENAI_API_KEY），不推荐作为长期生产主路径。",
    )

    add_heading(doc, "九、步骤 6：获取代码与环境变量", 1)
    add_para(
        doc,
        "将 Insight2Sale 私有仓库克隆到服务器某目录（或由我方提供 release 压缩包解压）。"
        "复制 .env.example 为 .env，按下一节填写。切勿将 .env 提交到 Git。",
    )

    add_heading(doc, "十、核心环境变量说明（.env）", 1)
    add_para(doc, "下表为生产环境最常见配置项，具体以仓库根目录 .env.example 注释为准。")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "变量名"
    hdr[1].text = "是否必填"
    hdr[2].text = "说明"
    for c in hdr:
        set_cell_shading(c, "E7E6E6")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Microsoft YaHei"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    rows = [
        ("DATABASE_URL", "必填", "PostgreSQL 连接串"),
        ("AUTH_SECRET", "必填", "随机长字符串，用于会话加密"),
        ("AUTH_URL", "必填", "对外访问根 URL，须与浏览器地址一致，含 https 与域名"),
        ("NEXTAUTH_URL", "必填", "建议与 AUTH_URL 相同"),
        ("ARK_API_KEY", "必填", "方舟 API Key"),
        ("ARK_BASE_URL", "必填", "一般为 https://ark.cn-beijing.volces.com/api/v3"),
        ("ARK_MODEL", "必填", "对话推理接入点 ID"),
        ("ARK_TIMEOUT_MS", "可选", "默认 8000 毫秒"),
        ("ARK_EMBEDDING_MODEL", "生产强烈建议", "Embedding 专用推理接入点 ID，与 ARK_MODEL 不同"),
        ("ARK_EMBEDDING_USE_MULTIMODAL", "视模型", "多模态 Embedding 时设为 1"),
        ("ARK_EMBEDDING_TIMEOUT_MS", "可选", "向量批量较慢时可加大"),
        ("NEXT_PUBLIC_SITE_URL", "可选", "仅当反向代理未传 Host 时作兜底"),
        ("VOLC_SPEECH_APP_KEY 等", "可选", "豆包语音 AppKey；另需 ACCESS_KEY、RESOURCE_ID（默认 volc.bigasr.auc_turbo）"),
        ("OPENAI_API_KEY", "可选", "Whisper 备选转写"),
    ]
    for name, req, desc in rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = req
        row[2].text = desc
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Microsoft YaHei"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    doc.add_paragraph()

    add_heading(doc, "十一、步骤 7：数据库迁移与种子数据", 1)
    add_para(doc, "在项目根目录依次执行（示例）：")
    cmds = [
        "npm ci 或 npm install",
        "npx prisma generate",
        "npx prisma migrate deploy   （生产推荐；将应用迁移到当前库）",
        "可选：npm run db:seed   （初始化演示账号与基础数据；生产是否执行由贵司决定）",
    ]
    for c in cmds:
        p = doc.add_paragraph(c, style="List Number")
        for run in p.runs:
            run.font.name = "Courier New"
    add_para(
        doc,
        "默认演示账号（若执行 seed）：admin、tianmanager、zhoulan、xuning；密码均为部署环境 .env 中的 DEFAULT_NEW_USER_PASSWORD（勿写入仓库）。"
        "首次使用默认密码登录后，系统会要求修改密码后方可进入工作台。",
    )

    add_heading(doc, "十二、步骤 8：知识库向量全量重算（启用 Embedding 后）", 1)
    add_para(
        doc,
        "在已配置 ARK_EMBEDDING_MODEL 且数据库中已有历史知识库切片的情况下，建议执行一次全量向量重算，避免新旧向量模型混用：",
    )
    add_para(doc, "npm run db:reembed-knowledge", bold=True)
    add_para(
        doc,
        "该命令会批量调用方舟 Embedding API，耗时与文档量、网络有关；请在业务低峰执行并观察控制台日志。",
    )

    add_heading(doc, "十三、步骤 9：构建与启动", 1)
    add_para(doc, "生产构建：npm run build")
    add_para(doc, "生产启动：npm run start（默认监听 3000；可通过环境变量 PORT 或启动参数指定端口，例如 3001）。")
    add_para(
        doc,
        "建议使用 PM2 或 systemd 守护进程，保证崩溃自启。示例（PM2）：pm2 start npm --name insight2sale -- start",
    )
    add_para(
        doc,
        "须在反向代理（Nginx）将 HTTPS 443 转发到本机 Node 端口；配置较大的 client_max_body_size 以支持知识库 PDF 上传（例如 32m 量级，与 next.config 中 serverActions 限制一致）。",
    )

    add_heading(doc, "十四、步骤 10：持久化目录（录音与上传）", 1)
    add_para(
        doc,
        "通话录音文件存储在服务器项目目录下 storage/call-recordings/；知识库等上传文件亦有 storage 下路径。"
        "请确保该目录在磁盘上持久化（容器部署时需挂载卷），并纳入贵司备份策略。不要将 storage 提交到 Git。",
    )

    add_heading(doc, "十五、步骤 11：验收清单（建议逐项打勾）", 1)
    checks = [
        "浏览器通过 https://贵司域名 可打开登录页，且与 AUTH_URL 一致。",
        "可使用账号登录并进入工作台（新账号完成强制改密流程）。",
        "知识库上传 PDF 后能完成向量化（已配置 ARK_EMBEDDING_MODEL）。",
        "解读台录音上传后，转写任务可完成（已配置 VOLC_SPEECH 或备选 Whisper）。",
        "服务器重启后服务自动恢复，录音与数据库数据不丢失。",
    ]
    for t in checks:
        p = doc.add_paragraph(t, style="List Bullet")
        for run in p.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    add_heading(doc, "十六、常见问题与排查", 1)
    faq = [
        ("登录失败或回调异常", "检查 AUTH_URL、NEXTAUTH_URL 是否与浏览器地址完全一致，含协议与端口。"),
        ("知识库检索结果异常或为空", "确认 ARK_EMBEDDING_MODEL 已配置且已执行重算；检查 embeddingModel 是否与当前接入点一致。"),
        ("向量请求超时", "尝试增大 ARK_EMBEDDING_TIMEOUT_MS 或 ARK_TIMEOUT_MS；检查服务器出网与方舟配额。"),
        ("录音无法转写", "检查 VOLC_SPEECH 密钥与资源 ID；查看 CallRecording 表中 processingError。"),
    ]
    for title, body in faq:
        add_para(doc, title + "：" + body, bold=False)

    add_heading(doc, "十七、安全与合规建议", 1)
    add_para(
        doc,
        "定期更新系统补丁；限制 SSH 与数据库暴露面；密钥仅保存在服务器权限受限文件中；"
        "生产数据库定期备份；对家长与销售数据按贵司合规要求留存与脱敏。",
    )

    add_heading(doc, "十八、可选：知识库 PDF OCR（扫描件 / 图片内文字）", 1)
    add_para(
        doc,
        "若知识库上传的 PDF 含扫描页或图片内文字，服务器需安装 Tesseract OCR（中英语言包）。"
        "可通过环境变量 PDF_OCR_DISABLE、PDF_OCR_MODE（hybrid / full）、PDF_OCR_MAX_PAGES 等调节行为，详见仓库根目录 .env.example 注释。",
    )

    add_heading(doc, "十九、附录：与内部设计文档的对应关系", 1)
    add_para(
        doc,
        "更细的产品与技术设计见交付仓库内 docs/design/system-design.md；"
        "通话录音见 §27；知识库向量见 §28。本文面向部署执行，不替代设计文档。",
    )

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("— 文档结束 —\nInsight2Sale v1.0.0")
    fr.font.size = Pt(9)
    fr.font.name = "Microsoft YaHei"
    fr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    doc.save(out)
    print(f"Written: {out}")


if __name__ == "__main__":
    main()
